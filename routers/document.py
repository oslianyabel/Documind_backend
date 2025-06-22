import asyncio
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from typing import Annotated

import asyncio
import aiofiles
import numpy
import PyPDF2
from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, HTTPException, UploadFile, status
from fastapi.responses import FileResponse
from sklearn.metrics.pairwise import cosine_similarity

from completions import Completions
from config import config, logger
from database import database, document_table, page_table, query_table
from models.document import (
    Document,
    DocumentWithSimilarity,
    PageWithSimilarity,
    SearchResult,
    UploadDocument,
    UserQuery,
)
from models.user import UserOut
from security import get_current_user
from utils import clean_text, get_embedding, get_page_embedding

router = APIRouter()
UserDep = Annotated[UserOut, Depends(get_current_user)]


@router.post("/upload", status_code=201)
async def upload_document(file: UploadFile, current_user: UserDep) -> UploadDocument:
    """
    Carga el documento y crea los embeddings en la base de datos
    """
    CHUNK_SIZE = 1024 * 1024
    process_list = []
    pages: list
    content = ""
    file_path = Path(config.DOCUMENT_PATH) / file.filename  # type: ignore
    async with aiofiles.open(file_path, "wb") as new_file:
        while chunk := await file.read(CHUNK_SIZE):
            await new_file.write(chunk)

    if file.content_type == "application/pdf":
        with open(file_path, "rb") as new_file:
            reader = PyPDF2.PdfReader(new_file)
            for idx, page in enumerate(reader.pages):
                content_page = clean_text(page.extract_text())
                if content_page:
                    process_list.append(get_page_embedding(idx, content_page))
                content += content_page + "\n"

            if process_list:
                pages = await asyncio.gather(*process_list)

    elif (
        file.content_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        doc = DocxDocument(file_path)  # type: ignore

        for idx, para in enumerate(doc.paragraphs):
            content_page = clean_text(para.text)
            if content_page:
                process_list.append(get_page_embedding(idx, content_page))
            content += content_page + "\n"

        if process_list:
            pages = await asyncio.gather(*process_list)

    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid Format File: {file.content_type} Must be .pdf or .docx",
        )

    data = {
        "name": file.filename,
        "url": f"{config.DOMAIN}/{config.DOCUMENT_PATH}/{file.filename}",
        "embeddings": await get_embedding(content[:2000]),
    }
    query = document_table.insert().values(data)
    id = await database.execute(query)

    for page in pages:  # type: ignore
        page["document_id"] = id
        query = page_table.insert().values(page)
        await database.execute(query)

    logger.info(f"Document {data['name']} with {len(pages)} pages was uploaded")  # type: ignore

    return UploadDocument(
        detail="file upload successfully",
        document_id=id,
        document_url=data["url"],
    )


@router.post("/search")
async def get_relevant_documents(
    user_query: UserQuery, limit: int = 3
) -> list[DocumentWithSimilarity]:
    user_query_embedding = await get_embedding(user_query.content)
    user_query_embedding = numpy.array(user_query_embedding).reshape(1, -1)

    query = document_table.select()
    documents = await database.fetch_all(query)

    def calculate_doc_similarity(doc) -> DocumentWithSimilarity | None:
        if doc.embeddings:  # type: ignore
            doc_embedding = numpy.array(doc.embeddings).reshape(1, -1)  # type: ignore
            similarity = cosine_similarity(user_query_embedding, doc_embedding)[0][0]

            return DocumentWithSimilarity(
                id=doc.id,  # type: ignore
                name=doc.name,  # type: ignore
                url=doc.url,  # type: ignore
                similarity=float(similarity),
            )

        return None

    with ThreadPoolExecutor() as executor:
        documents_with_similarity = list(
            executor.map(calculate_doc_similarity, documents)
        )

    documents_with_similarity = [doc for doc in documents_with_similarity if doc]
    documents_with_similarity.sort(key=lambda doc: doc.similarity, reverse=True)

    return documents_with_similarity[:limit]


@router.get("/{document_id}", response_model=Document)
async def get_document(document_id: int):
    query = document_table.select().where(document_table.c.id == document_id)
    document = await database.fetch_one(query)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    return document


@router.get("/{document_id}/download")
async def download_document(document_id: int) -> FileResponse:
    query = document_table.select().where(document_table.c.id == document_id)
    document = await database.fetch_one(query)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    path = document.url.replace(config.DOMAIN, "")  # type: ignore
    if path.startswith("/"):
        path = path[1:]

    print(path)
    file_path = Path(path)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found on server")

    return FileResponse(
        file_path,
        filename=document.name,  # type: ignore
        headers={"Content-Disposition": f"attachment; filename={document.name}"},  # type: ignore
    )


@router.post("/{document_id}/search")
async def get_document_response(
    document_id: int, user_query: UserQuery, limit: int = 3
) -> list[PageWithSimilarity]:
    query = document_table.select().where(document_table.c.id == document_id)
    document = await database.fetch_one(query)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    user_query_embedding = await get_embedding(user_query.content)
    user_query_embedding = numpy.array(user_query_embedding).reshape(1, -1)

    query = page_table.select().where(page_table.c.document_id == document.id)  # type: ignore
    pages = await database.fetch_all(query)

    def calculate_page_similarity(page) -> PageWithSimilarity | None:
        if page.embeddings:  # type: ignore
            page_embedding = numpy.array(page.embeddings).reshape(1, -1)  # type: ignore
            similarity = cosine_similarity(user_query_embedding, page_embedding)[0][0]

            return PageWithSimilarity(
                document_id=page.document_id,  # type: ignore
                page_number=page.page_number,  # type: ignore
                similarity=float(similarity),
            )

        return None

    with ThreadPoolExecutor() as executor:
        pages_with_similarity = list(executor.map(calculate_page_similarity, pages))

    pages_with_similarity = [page for page in pages_with_similarity if page]
    pages_with_similarity.sort(key=lambda page: page.similarity, reverse=True)

    return pages_with_similarity[:limit]


@router.post("/{document_id}/page/{page_number}/search")
async def get_page_response(
    document_id: int, page_number: int, user_query: UserQuery
) -> SearchResult:
    query = (
        page_table.select()
        .where(page_table.c.document_id == document_id)
        .where(page_table.c.page_number == page_number)
    )
    page = await database.fetch_one(query)

    messages = [
        {
            "role": "system",
            "content": f"El usuario te hará una consulta que debes responder con contenido LITERAL tomado del siguiente texto: {page.content}. Si la respuesta no aparece tu respuesta será: No Answer",  # type: ignore
        },
        {"role": "user", "content": user_query.content},  # type: ignore
    ]
    ans = Completions().submit_message(messages)

    query_data = {
        "query": user_query.content,
        "answer": ans,
        "document_id": document_id,
        "page": page_number,
    }
    query = query_table.insert().values(query_data)
    await database.execute(query)

    return SearchResult(**query_data)
